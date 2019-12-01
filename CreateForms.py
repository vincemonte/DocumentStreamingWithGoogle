from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
import gspread
from oauth2client.service_account import ServiceAccountCredentials
#Author: vincemonte 
# v 1.0

'''
Setting up the credentials for spreadsheet access. Please note that
the .json file used to provide authentication through the oauth2client module
hasn't been included for my own safety.
'''
scope = ['https://www.googleapis.com/auth/spreadsheets',
"https://www.googleapis.com/auth/drive.file"
,"https://www.googleapis.com/auth/drive"]
creds = ServiceAccountCredentials.from_json_keyfile_name("creds.json", scope)
client = gspread.authorize(creds)

'''
Gathers all data from the Google Spreadsheet and transforms the data into
a different format
Data will be stored in a matrix (same as spreadsheet), instead a list of lists
*Note: each row is stored as a list containing a dictionary by Google's default*
'''
def gatherInfoFromSpreadsheets():
    allInfoList = []
    sheet = client.open("Testing of Incident Response Form (Responses)").sheet1
    data = sheet.get_all_records()
    for row in data:
        rowData = []
        for value in row.values():
            rowData.append(value)
        allInfoList.append(rowData)
    return allInfoList

'''
Creates a single form document, utilizing the helperRun and helperHeading
functions implemented below. Note that this isn't the most efficient way
to utilize the docx library as paragraphs can be represented as objects and
thus a format/style may be set for the entire paragraph. However, adjusting to
implement the different styles all present in one paragraph required a different
and less efficient approach.
'''
def createForm(singleEntry):
    doc = Document()
    doc.add_picture('helios-header-try3.png', width = Inches(6),
    height=Inches(1.6))
    #Gathering to label docs
    timeStamp = singleEntry[0]
    timeStamp = timeStamp[10:]
    timeStamp = '({})'.format(timeStamp)
    #helperHeading(doc, timeStamp, 3, 'E6582D','Calibri') #don't use helper here
    doc.add_paragraph('Reported by: {}'.format(singleEntry[1])).underline=True
    #Incident Description Section
    helperHeading(doc,'Incident Description', 'E6582D', 'Calibri')
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    helperRun(p, 'Time of First Attack: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[2], '1C1006', 1, 1)
    helperRun(p, 'Time of Attack Detection: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[3] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Estimated Recovery Time: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[4], '1C1006', 1, 1)
    helperRun(p, 'Duration of Attack: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[5] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Has the Attack Ended?: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[6], '1C1006', 1, 1)
    helperRun(p, 'Estimated Damage Account: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[7] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Number of Hosts Affected: ', 'EB8232', 0, 0)
    helperRun(p, str(singleEntry[8]), '1C1006', 1, 1)
    helperRun(p, 'Number of Users Affected: ', 'EB8232', 0, 0)
    helperRun(p, str(singleEntry[9]) + '\n', '1C1006', 1, 0)
    helperRun(p, 'Severity of Attack: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[10], '1C1006', 1, 1)
    helperRun(p, 'Attack Category: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[11] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Type of Incident: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[12] + '\n', '1C1006', 1, 0)
    helperRun(p, 'How Did You Become Initially Aware of the Attack?\n', 'EB8232', 0, 0)
    helperRun(p, singleEntry[13], '1C1006', 0, 0)
    #Impact of Attack Section
    helperHeading(doc, 'Impact of Attack', 'E6582D', 'Calibri')
    p=doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    helperRun(p, 'Primary Purpose of Host: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[14] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Host Name: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[15] , '1C1006', 1, 1)
    helperRun(p, 'IP Address: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[16] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Operating System: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[17] + '\n', '1C1006', 1, 0)
    helperRun(p, 'Applications Affected: ', 'EB8232', 0, 0)
    helperRun(p, singleEntry[18], '1C1006', 1, 0)
    #Data Loss/Compromise Section
    helperHeading(doc, 'Data Loss/Compromise', 'E6582D', 'Calibri')
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    helperRun(p, 'Was there any loss or compromise of sensitive data?:\n', 'EB8232', 0, 0)
    helperRun(p, singleEntry[19] + '\n', '1C1006', 0, 0)
    helperRun(p, 'Was there any change to the afflicted system?:\n', 'EB8232', 0, 0)
    helperRun(p, singleEntry[20] + '\n', '1C1006', 0, 0)
    helperRun(p, 'Was the HMI/ICS affected?:\n', 'EB8232', 0, 0)
    helperRun(p, singleEntry[21] + '\n', '1C1006', 0, 0)
    helperRun(p, 'Corrective measures enacted:\n', 'EB8232', 0, 0)
    helperRun(p, singleEntry[22] + '\n', '1C1006', 0, 0)
    #saving the document
    doc.save('IncidentResponseForm{}.docx'.format(timeStamp))


'''
Utilizing the gatherInfoFromSpreadsheets function, this function passses in a
row (represented as a list) into an individual createForm call. This allows
for a .docx file to be made for each entry all on one function call.
'''
def createAllWordForms():
    allEntries = gatherInfoFromSpreadsheets()
    for entry in allEntries:
        createForm(entry)

'''
Helper function to allow for the easy insertion of a custom heading in the
.docx file
'''
def helperHeading(doc, text, hexColor, font_name):
    styles = doc.styles
    new_heading_style = styles.add_style('New Heading{}'.format(text), WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.name = font_name
    hex_tuple = RGBColor.from_string(hexColor)
    font.color.rgb = RGBColor(hex_tuple[0], hex_tuple[1], hex_tuple[2])
    #font.size = Pt(font_size)
    doc.add_paragraph(text, style='New Heading{}'.format(text))
'''
Helper function to allow for the easy insertion of a custom
run statement in the .docx file.
'''
def helperRun(paragraph, text, hexColor, underline, tab):
    run = paragraph.add_run(text)
    run.underline = bool(underline)
    font = run.font
    #font.size = Pt(font_size)
    hex_tuple = RGBColor.from_string(hexColor)
    font.color.rgb = RGBColor(hex_tuple[0], hex_tuple[1], hex_tuple[2])
    font.name = 'Athelas'
    if bool(tab) is True:
        paragraph.add_run().add_tab()
        paragraph.add_run().add_tab()

'''
This function is responsible to create a text document with all the data
collected. I included this so if any member needed to review any incident,
I could quickly send a text file instead of seperate .docx files.
'''
def createTextFileForm():
    sheet = client.open("Testing of Incident Response Form (Responses)").sheet1
    data = sheet.get_all_records()
    outFile = open('All_Responses.txt', 'w')
    for row in data:
        for key in row.keys():
            outFile.write('{}: {}\n'.format(key, row.get(key)))
        outFile.write('\n\n')
    outFile.close()

createTextFileForm()
createAllWordForms()
